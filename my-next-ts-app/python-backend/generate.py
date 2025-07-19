from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from docx import Document
import uuid
import os
import re

app = FastAPI()

@app.get("/")
def read_root():
    return {"message": "FastAPI server is running"}

@app.get("/generate-doc")
def get_generate_doc():
    return {"message": "Use POST method to generate document"}

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://127.0.0.1:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class FormData(BaseModel):
    refType: str = ""
    refNumber: str = ""
    appType: str = ""
    applicantName: str = ""
    gender: str = ""
    nationality: str = ""
    country: str = ""
    age: str = ""
    address: str = ""
    email: str = ""
    contact: str = ""
    applicantCategory: str = ""
    inventorSame: str = ""
    inventorName: str = ""
    inventorGender: str = ""
    inventorNationality: str = ""
    inventorCountry: str = ""
    inventorAge: str = ""
    inventorAddress: str = ""
    inventorEmail: str = ""
    inventorContact: str = ""
    title: str = ""
    agentNumber: str = ""
    agentName: str = ""
    agentMobile: str = ""
    serviceName: str = ""
    serviceAddress: str = ""
    servicePhone: str = ""
    serviceMobile: str = ""
    serviceFax: str = ""
    serviceEmail: str = ""
    convCountry: str = ""
    convAppNumber: str = ""
    convFilingDate: str = ""
    convApplicant: str = ""
    convTitle: str = ""
    convIPC: str = ""
    pctNumber: str = ""
    pctDate: str = ""
    divAppNumber: str = ""
    divDate: str = ""
    mainAppNumber: str = ""
    mainDate: str = ""
    specPages: str = ""
    claims: str = ""
    abstract: str = ""
    drawings: str = ""
    totalFees: str = ""
    paymentMethod: str = ""
    paymentNumber: str = ""
    paymentDate: str = ""
    paymentBank: str = ""
    submissionDate: str = ""
    applicantNameFinal: str = ""

def replace_text_in_paragraph(paragraph, replacements):
    """
    Replace text in a paragraph while preserving formatting.
    This handles cases where placeholders might be split across multiple runs.
    """
    # Get the full text of the paragraph
    full_text = paragraph.text
    
    # Check if any placeholder exists in this paragraph
    has_placeholder = any(placeholder in full_text for placeholder in replacements.keys())
    
    if has_placeholder:
        # Replace all placeholders in the full text
        new_text = full_text
        for placeholder, replacement in replacements.items():
            new_text = new_text.replace(placeholder, replacement)
        
        # If text changed, update the paragraph
        if new_text != full_text:
            # Clear existing runs
            for run in paragraph.runs:
                run.clear()
            
            # Add the new text as a single run
            paragraph.add_run(new_text)

def replace_text_in_run(run, replacements):
    """
    Replace text in individual runs.
    """
    for placeholder, replacement in replacements.items():
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, replacement)

def process_table_cell(cell, replacements):
    """
    Process each cell in a table, handling nested tables and all paragraphs.
    """
    # Process all paragraphs in the cell
    for paragraph in cell.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)
        # Also process individual runs for additional safety
        for run in paragraph.runs:
            replace_text_in_run(run, replacements)
    
    # Handle nested tables within cells
    for nested_table in cell.tables:
        process_table(nested_table, replacements)

def process_table(table, replacements):
    """
    Process entire table including all rows, cells, and nested structures.
    """
    for row in table.rows:
        for cell in row.cells:
            process_table_cell(cell, replacements)

def replace_in_headers_footers(doc, replacements):
    """
    Replace text in headers and footers.
    """
    # Process headers
    for section in doc.sections:
        # Primary header
        if section.header:
            for paragraph in section.header.paragraphs:
                replace_text_in_paragraph(paragraph, replacements)
                for run in paragraph.runs:
                    replace_text_in_run(run, replacements)
            
            # Process tables in headers
            for table in section.header.tables:
                process_table(table, replacements)
        
        # Primary footer
        if section.footer:
            for paragraph in section.footer.paragraphs:
                replace_text_in_paragraph(paragraph, replacements)
                for run in paragraph.runs:
                    replace_text_in_run(run, replacements)
            
            # Process tables in footers
            for table in section.footer.tables:
                process_table(table, replacements)

def comprehensive_text_replacement(doc, replacements):
    """
    Perform comprehensive text replacement throughout the entire document.
    """
    # Process main document paragraphs
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)
        for run in paragraph.runs:
            replace_text_in_run(run, replacements)
    
    # Process all tables in the main document
    for table in doc.tables:
        process_table(table, replacements)
    
    # Process headers and footers
    replace_in_headers_footers(doc, replacements)

@app.post("/generate-doc")
def generate_doc(data: FormData):
    template_path = "ui/Form1.docx"
    output_dir = "generated"
    os.makedirs(output_dir, exist_ok=True)
    output_path = f"{output_dir}/IPR_Form_{uuid.uuid4().hex[:8]}.docx"

    doc = Document(template_path)

    replacements = {
        "[refType]": data.refType,
        "[refNumber]": data.refNumber,
        "[appType]": data.appType,
        "[applicantName]": data.applicantName,
        "[gender]": data.gender,
        "[nationality]": data.nationality,
        "[country]": data.country,
        "[age]": data.age,
        "[address]": data.address,
        "[email]": data.email,
        "[contact]": data.contact,
        "[applicantCategory]": data.applicantCategory,
        "[inventorSame]": data.inventorSame,
        "[inventorName]": data.inventorName,
        "[inventorGender]": data.inventorGender,
        "[inventorNationality]": data.inventorNationality,
        "[inventorCountry]": data.inventorCountry,
        "[inventorAge]": data.inventorAge,
        "[inventorAddress]": data.inventorAddress,
        "[inventorEmail]": data.inventorEmail,
        "[inventorContact]": data.inventorContact,
        "[title]": data.title,
        "[agentNumber]": data.agentNumber,
        "[agentName]": data.agentName,
        "[agentMobile]": data.agentMobile,
        "[serviceName]": data.serviceName,
        "[serviceAddress]": data.serviceAddress,
        "[servicePhone]": data.servicePhone,
        "[serviceMobile]": data.serviceMobile,
        "[serviceFax]": data.serviceFax,
        "[serviceEmail]": data.serviceEmail,
        "[convCountry]": data.convCountry,
        "[convAppNumber]": data.convAppNumber,
        "[convFilingDate]": data.convFilingDate,
        "[convApplicant]": data.convApplicant,
        "[convTitle]": data.convTitle,
        "[convIPC]": data.convIPC,
        "[pctNumber]": data.pctNumber,
        "[pctDate]": data.pctDate,
        "[divAppNumber]": data.divAppNumber,
        "[divDate]": data.divDate,
        "[mainAppNumber]": data.mainAppNumber,
        "[mainDate]": data.mainDate,
        "[specPages]": data.specPages,
        "[claims]": data.claims,
        "[abstract]": data.abstract,
        "[drawings]": data.drawings,
        "[totalFees]": data.totalFees,
        "[paymentMethod]": data.paymentMethod,
        "[paymentNumber]": data.paymentNumber,
        "[paymentDate]": data.paymentDate,
        "[paymentBank]": data.paymentBank,
        "[submissionDate]": data.submissionDate,
        "[applicantNameFinal]": data.applicantNameFinal,
    }

    # Use the comprehensive replacement function
    comprehensive_text_replacement(doc, replacements)

    doc.save(output_path)

    return FileResponse(
        path=output_path,
        filename="IPR_Form1.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )